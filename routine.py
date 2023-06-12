from openpyxl import Workbook
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle

# Define your routine with specific time slots
routine = {
    "Morning Routine": [
        {"time": "6:00 AM", "task": "Wake up early and perform Fajr prayer."},
        {"time": "6:30 AM", "task": "Spend 30 minutes on English language practice."},
        {"time": "7:00 AM", "task": "Take a shower and get ready for the day."},
        {"time": "7:30 AM", "task": "Have a healthy breakfast."},
        {"time": "8:00 AM", "task": "Competitive programming practice"},
        {"time": "9:00 AM", "task": "Machine learning course"},
        {"time": "10:00 AM", "task": "Backend course"}
    ],
    "University Days (Monday to Thursday)": [
        {"time": "9:00 AM", "task": "Attend university."},
        {"time": "1:00 PM", "task": "Review class notes or work on assignments during breaks."},
        {"time": "5:00 PM", "task": "Make a to-do list and prioritize tasks."},
        {"time": "7:00 PM", "task": "Allocate time for studying."},
        {"time": "8:00 PM", "task": "Codeforces competition"}
    ],
    "Evening Routine": [
        {"time": "5:00 PM", "task": "Perform Asr prayer."},
        {"time": "6:00 PM", "task": "Allocate time for learning and practicing programming languages."},
        {"time": "7:00 PM", "task": "Engage in physical exercise or recreational activities."},
        {"time": "8:00 PM", "task": "Spend quality time with family or friends."},
        {"time": "9:00 PM", "task": "Perform Maghrib prayer."},
        {"time": "9:30 PM", "task": "Allocate time for self-improvement or learning new concepts."},
        {"time": "10:30 PM", "task": "Perform Isha prayer."},
        {"time": "11:00 PM", "task": "Relax and engage in hobbies."},
        {"time": "11:30 PM", "task": "Get 7-8 hours of sleep each night."}
    ],
    "Weekends (Friday and Saturday)": [
        {"time": "9:00 AM", "task": "Dedicate time to studying and working on assignments."},
        {"time": "1:00 PM", "task": "Review topics covered during the week and clarify doubts."},
        {"time": "3:00 PM", "task": "Spend more time on learning programming languages and practicing coding exercises."},
        {"time": "5:00 PM", "task": "Take breaks and engage in enjoyable activities."},
        {"time": "7:00 PM", "task": "Use time for personal development and improving English speaking skills."},
        {"time": "8:00 PM", "task": "Competitive programming course"}
    ]
}


# Generate PDF
pdf_filename = 'routine.pdf'

# Create a PDF document
doc = SimpleDocTemplate(pdf_filename, pagesize=letter)

# Create a table from routine data
table_data = [['Time', 'Task']]
for section, tasks in routine.items():
    table_data.append([section, ''])
    for task in tasks:
        table_data.append([task['time'], task['task']])

# Create the table and apply styles
table = Table(table_data)
table.setStyle(TableStyle([
    ('BACKGROUND', (0, 0), (-1, 0), '#4CAF50'),  # Header background color
    ('TEXTCOLOR', (0, 0), (-1, 0), '#FFFFFF'),  # Header text color
    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),  # Center alignment for all cells
    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),  # Header font
    ('FONTSIZE', (0, 0), (-1, 0), 12),  # Header font size
    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),  # Header bottom padding
    ('BACKGROUND', (0, 1), (-1, -1), '#f2f2f2'),  # Content background color
    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),  # Content font
    ('FONTSIZE', (0, 1), (-1, -1), 10),  # Content font size
    ('BOTTOMPADDING', (0, 1), (-1, -1), 6),  # Content bottom padding
]))

# Build the PDF document
elements = [table]
doc.build(elements)

# Generate XLSX
wb = Workbook()
ws = wb.active
ws.title = 'Routine'
row = 1
for section, tasks in routine.items():
    ws.cell(row=row, column=1, value=section)
    row += 1
    for task in tasks:
        ws.cell(row=row, column=2, value=task["time"])
        ws.cell(row=row, column=3, value=task["task"])
        row += 1
wb.save('routine.xlsx')

# Generate Word
document = Document()
for section, tasks in routine.items():
    document.add_heading(section, level=1)
    for task in tasks:
        document.add_paragraph(f"{task['time']}: {task['task']}")
document.save('routine.docx')
