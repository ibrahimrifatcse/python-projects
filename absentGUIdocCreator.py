# please do install this module : pip install python-docx
#  pip install pySimpleGUI

import string
import docx
import PySimpleGUI as sg

# Create the GUI layout
layout = [
    [sg.Text("All student IDs file:"), sg.Input(key="all_ids_file"), sg.FileBrowse()],
    [sg.Text("Garbage IDs file:"), sg.Input(key="garbage_ids_file"), sg.FileBrowse()],
    [sg.Button("Run"), sg.Button("Exit")]
]

# Create the window
window = sg.Window("Absent Present Students", layout)

# Event loop to process user input
while True:
    event, values = window.read()
    if event == "Run":
        all_ids_file = values["all_ids_file"]
        garbage_ids_file = values["garbage_ids_file"]
        
        # Open files for reading
        try:
            all_ids_file = open(all_ids_file, "r")
            garbage_ids_file = open(garbage_ids_file, "r")
        except IOError:
            sg.popup_error("Error: Unable to open file.")
            continue
        
        # Read all IDs from all_ids.txt into a list
        all_ids = []
        for line in all_ids_file:
            all_ids.append(line.strip())

        all_ids_file.close()

        # Read all IDs from garbage_ids.txt into a list
        garbage_ids = []
        for line in garbage_ids_file:
            # Extract 9-digit student IDs from the line
            ids_in_line = [id.translate(str.maketrans('', '', string.punctuation)) for id in line.split() if len(id) == 9 and id.isdigit()]
            garbage_ids += ids_in_line

        garbage_ids_file.close()

        # Find the missing IDs
        missing_ids = set(all_ids) - set(garbage_ids)

        # Create a new text file for absent students
        with open("AbsentStudents.txt", "w") as absent_file:
            for student_id in sorted(missing_ids):
                absent_file.write(student_id + "\n")

        # Create a new Word document for present students
        document = docx.Document()
        document.add_heading("Present Students", level=0)

        for i, student_id in enumerate(sorted(set(all_ids) - missing_ids), 1):
            document.add_paragraph(f"{i}. {student_id}")

        document.save("PresentStudents.docx")
        
        sg.popup("Operation complete!")
    elif event == "Exit":
        break

# Close the window
window.close()
